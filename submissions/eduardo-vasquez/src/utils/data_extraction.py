from typing import List, Dict
from docx import Document
import io
import re
from typing import Tuple
from utils.clients import create_google_storage_client
from utils.custom_logging import get_logger

logger = get_logger()


def download_docx_from_gcs(bucket_name: str) -> List:
    """
    Downloads DOCX files from a Google Cloud Storage bucket.

    Args:
        bucket_name (str): Name of the Google Cloud Storage bucket.

    Returns:
        List: A list of DOCX blobs (GCS files) from the bucket, or an empty list in case of an error.
    """
    try:
        # Create a Google Cloud Storage client
        storage_client = create_google_storage_client()

        # Get the bucket object
        logger.info(f"Getting bucket '{bucket_name}'...")
        bucket = storage_client.bucket(bucket_name)

        # List all blobs (files) in the bucket
        logger.info("Getting blobs from bucket...")
        blobs = list(bucket.list_blobs())

        # Filter DOCX files (optional)
        docx_blobs = [blob for blob in blobs if blob.name.endswith(".docx")]

        return docx_blobs

    except Exception as e:
        logger.error(f"Error downloading DOCX files from GCS: {e}")


def extract_date_from_docx(doc: Document) -> str:
    """
    Extracts the date from a DOCX document.

    Args:
        doc (Document): The DOCX document object.

    Returns:
        str: The extracted date or None if not found.
    """
    pattern_date = r"Date:\s*(?P<date>\d{2}/\d{2}/\d{4})"  # Format like '04/05/2025'
    try:
        for para in doc.paragraphs:
            match_date = re.search(pattern_date, para.text)
            if match_date:
                return match_date.group("date")  # Return the first found date
    except Exception as e:
        logger.error(f"Error extracting date: {e}", exc_info=True)


def process_paragraphs(doc: Document, date: str, blob: str) -> List[Dict[str, object]]:
    """
    Processes paragraphs in a DOCX document, extracts metadata, and stores the text.

    Args:
        doc (Document): The DOCX document.
        pattern_transcription (str): The regex pattern for extracting transcription data.
        date (str): The extracted meeting date.
        blob (str): The GCS blob name (filename).

    Returns:
        List[Dict[str, object]]: A list of dictionaries containing text and metadata.
    """
    extracted_data = []
    message_counter = 0  # Message number counter
    current_speaker = None
    meeting_title = blob.split(".")[0]

    # Regex patterns
    pattern_transcription = re.compile(
        r"^\[(?P<timestamp>\d{2}:\d{2}:\d{2})\]\s*(?P<speaker_name>[^(]+?)\s*\((?P<position>[^)]+)\):\s*$"
    )

    try:
        for para in doc.paragraphs:
            para_text = para.text.strip()

            # Match speaker lines
            if speaker_match := pattern_transcription.match(para_text):
                current_speaker = {
                    key: speaker_match.group(key).strip()
                    for key in ["timestamp", "speaker_name", "position"]
                }
                continue

            # Extract speaker's message
            if (
                current_speaker
                and para_text.startswith('"')
                and para_text.endswith('"')
            ):
                message_counter += 1
                extracted_data.append(
                    {
                        "text": para_text[1:-1],  # Remove quotes around text
                        "name": meeting_title,
                        "metadata": {
                            "date": date,
                            **current_speaker,
                            "message_number": message_counter,
                        },
                    }
                )
                current_speaker = None  # Reset current speaker

        return extracted_data
    except Exception as e:
        logger.error(f"Unexpected error while processing paragraphs: {e}")


def parse_docx_file(blob: str, file_bytes: bytes) -> List[Dict[str, object]]:
    """
    Parses a DOCX file to extract text and metadata.

    Args:
        blob (str): The GCS blob name (filename).
        file_bytes (bytes): The byte content of the DOCX file.

    Returns:
        List[Dict[str, object]]: A list of dictionaries containing text and metadata.
    """

    # Read DOCX file in memory
    with io.BytesIO(file_bytes) as docx_file:
        doc = Document(docx_file)

        # Extract the date (assuming it's in the first few paragraphs)
        date = extract_date_from_docx(doc)

        # Process each paragraph in the DOCX file
        extracted_data = process_paragraphs(doc, date, blob)

    return extracted_data


def extract_text_and_metadata(
    bucket_name: str,
    extract_all: bool = False,
) -> Tuple[List[Dict[str, object]], str]:
    """
    Main function to extract text and metadata from DOCX files in a GCS bucket.

    Args:
        bucket_name (str): Name of the Google Cloud Storage bucket.
        extract_all (bool): If True, extracts all available DOCX files in the bucket.
                            If False, extracts only the latest DOCX.

    Returns:
        Tuple[List[Dict[str, object]], str]:
            - A list of dictionaries containing metadata (if extract_metadata=True).
            - A single string containing the full extracted text from the document.
    """
    # Step 1: Download DOCX files from GCS
    blobs = download_docx_from_gcs(bucket_name)

    extracted_data = []

    try:
        # Select blobs based on the extract_all flag
        selected_blobs = (
            blobs if extract_all else [max(blobs, key=lambda blob: blob.updated)]
        )

        for blob in selected_blobs:
            logger.info(f"Processing blob: {blob}")
            file_bytes = blob.download_as_bytes()

            logger.info(f"Processing file_bytes: {file_bytes}")
            file_data = parse_docx_file(blob.name, file_bytes)
            logger.info(f"Processing file_data: {file_data}")
            extracted_data.extend(file_data)

            logger.info("Extracting document...")
            with io.BytesIO(file_bytes) as docx_file:
                doc = Document(docx_file)
                full_document = "\n".join([para.text for para in doc.paragraphs])

        return extracted_data, full_document

    except Exception as e:
        logger.error(f"Error processing {blob.name}: {e}")
